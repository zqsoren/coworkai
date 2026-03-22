from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks
from pydantic import BaseModel
from typing import Optional, List, Dict
import os
import shutil

from src.utils.rag_ingestion import RAGIngestion
from src.core.file_manager import FileManager

router = APIRouter(prefix="/api/knowledge", tags=["knowledge"])

# Models
class ProcessRequest(BaseModel):
    workspace_id: str
    agent_id: str

class FileListResponse(BaseModel):
    files: List[str]

# Context
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DATA_ROOT = os.path.join(PROJECT_ROOT, "data")
file_manager = FileManager(DATA_ROOT)

@router.get("/files", response_model=FileListResponse)
def list_files(workspace_id: str, agent_id: str, type: str = "knowledge_base"):
    """
    List files in a specific directory.
    type options: 
      - 'knowledge_base/uploads'
      - 'knowledge_base/processed'
      - 'context/static', 'context/active', 'context/archives'
    """
    # Security: Ensure type doesn't contain traversal (simple check)
    if ".." in type:
        raise HTTPException(status_code=400, detail="Invalid path")
    
    dir_path = os.path.join(DATA_ROOT, workspace_id, agent_id, type)
    
    if not os.path.exists(dir_path):
        return {"files": []}
        
    try:
        files = [f for f in os.listdir(dir_path) if os.path.isfile(os.path.join(dir_path, f))]
        return {"files": sorted(files)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload")
def upload_file(
    workspace_id: str = Form(...),
    agent_id: str = Form(...),
    type: str = Form(...), # 'knowledge_base' (auto-maps to uploads), or direct paths
    files: List[UploadFile] = File(...)
):
    """Upload files to the specified directory."""
    # Special handling for KB upload -> defaults to 'uploads' folder
    if type == "knowledge_base":
        target_sub = "knowledge_base/uploads"
    elif type == "chat_upload":
        target_sub = "shared/uploads"
    else:
        target_sub = type

    target_dir = os.path.join(DATA_ROOT, workspace_id, agent_id, target_sub)
    os.makedirs(target_dir, exist_ok=True)
    
    saved_files = []
    extracted_texts = {}  # chat_upload 时返回提取的文本

    try:
        for file in files:
            file_path = os.path.join(target_dir, file.filename)
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            saved_files.append(file.filename)

            # chat_upload 模式：提取文本内容后删除临时文件
            if type == "chat_upload":
                try:
                    ext = os.path.splitext(file.filename)[1].lower()
                    text = ""
                    print(f"[ChatUpload] Extracting text from: {file.filename}, ext={ext}, path={file_path}")
                    if ext in (".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".xml", ".html", ".py", ".js", ".ts", ".css"):
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                        print(f"[ChatUpload] Text file read OK, length={len(text)}")
                    elif ext in (".doc", ".docx"):
                        try:
                            import docx
                            doc = docx.Document(file_path)
                            parts = []
                            # 1. Body paragraphs
                            for p in doc.paragraphs:
                                if p.text.strip():
                                    parts.append(p.text)
                            # 2. Tables (many resumes use table layout)
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                                    if row_text:
                                        parts.append(" | ".join(row_text))
                            # 3. Headers and Footers
                            for section in doc.sections:
                                if section.header and section.header.paragraphs:
                                    for p in section.header.paragraphs:
                                        if p.text.strip():
                                            parts.append(p.text)
                                if section.footer and section.footer.paragraphs:
                                    for p in section.footer.paragraphs:
                                        if p.text.strip():
                                            parts.append(p.text)
                            text = "\n".join(parts)
                            print(f"[ChatUpload] DOCX read OK, paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}, text_length={len(text)}")
                        except Exception as docx_err:
                            print(f"[ChatUpload] DOCX read FAILED: {type(docx_err).__name__}: {docx_err}")
                            text = f"[无法解析 {ext} 文件: {docx_err}]"
                    elif ext == ".pdf":
                        try:
                            from PyPDF2 import PdfReader
                            reader = PdfReader(file_path)
                            pages = []
                            for page in reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    pages.append(page_text)
                            text = "\n".join(pages)
                            print(f"[ChatUpload] PDF read OK, pages={len(reader.pages)}, text_length={len(text)}")
                        except Exception as pdf_err:
                            print(f"[ChatUpload] PDF read FAILED: {type(pdf_err).__name__}: {pdf_err}")
                            text = f"[无法解析 PDF: {pdf_err}]"
                    else:
                        # 尝试以文本方式读取未知格式
                        try:
                            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                                text = f.read()
                        except Exception:
                            text = f"[不支持的文件格式: {ext}]"
                    extracted_texts[file.filename] = text.strip()
                    print(f"[ChatUpload] Final extracted text length for {file.filename}: {len(extracted_texts[file.filename])}")
                except Exception as e:
                    import traceback
                    print(f"[ChatUpload] EXCEPTION during extraction: {traceback.format_exc()}")
                    extracted_texts[file.filename] = f"[文件读取失败: {e}]"
                finally:
                    # 删除临时文件，chat_upload 不需要持久保存
                    try:
                        os.remove(file_path)
                    except Exception:
                        pass

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

    result = {"status": "success", "saved": saved_files}
    if extracted_texts:
        result["extracted_texts"] = extracted_texts
    return result

@router.delete("/file")
def delete_file(workspace_id: str, agent_id: str, type: str, filename: str):
    """Delete a file."""
    target_path = os.path.join(DATA_ROOT, workspace_id, agent_id, type, filename)
    
    # Security check: ensure strictly within data root
    resolved = os.path.abspath(target_path)
    if not resolved.startswith(os.path.abspath(DATA_ROOT)):
         raise HTTPException(status_code=403, detail="Access denied")
         
    if os.path.exists(resolved):
        os.remove(resolved)
        return {"status": "success"}
    else:
        raise HTTPException(status_code=404, detail="File not found")

@router.post("/process")
def process_knowledge_base(request: ProcessRequest):
    """
    Trigger RAG ingestion:
    1. Scan knowledge_base/uploads/
    2. Ingest valid files (split + embed)
    3. Move to knowledge_base/processed/
    """
    import traceback
    
    print(f"[Knowledge] Processing request: workspace={request.workspace_id}, agent={request.agent_id}")
    
    try:
        uploads_dir = os.path.join(DATA_ROOT, request.workspace_id, request.agent_id, "knowledge_base/uploads")
        processed_dir = os.path.join(DATA_ROOT, request.workspace_id, request.agent_id, "knowledge_base/processed")
        
        print(f"[Knowledge] uploads_dir: {uploads_dir}")
        print(f"[Knowledge] uploads_dir exists: {os.path.exists(uploads_dir)}")
        
        if not os.path.exists(uploads_dir):
            return {"status": "success", "results": {}, "message": "No uploads found"}
            
        os.makedirs(processed_dir, exist_ok=True)
        
        ingestion = RAGIngestion(DATA_ROOT, request.workspace_id, request.agent_id)
        results = {}
        
        files_in_dir = [f for f in os.listdir(uploads_dir) if os.path.isfile(os.path.join(uploads_dir, f))]
        print(f"[Knowledge] Files to process: {files_in_dir}")
        
        for filename in files_in_dir:
            src_path = os.path.join(uploads_dir, filename)
            print(f"[Knowledge] Processing file: {filename}")
                
            try:
                # Ingest: Load → Clean → Split → Embed → Store
                count = ingestion.ingest_file(src_path)
                results[filename] = {"chunks": count, "status": "success"}
                print(f"[Knowledge] ✓ {filename}: {count} chunks ingested")
                
                # Move to processed
                dst_path = os.path.join(processed_dir, filename)
                if os.path.exists(dst_path):
                    os.remove(dst_path)
                shutil.move(src_path, dst_path)
                print(f"[Knowledge] ✓ {filename}: moved to processed")
            except Exception as e:
                error_msg = str(e)
                print(f"[Knowledge] ✗ {filename}: Error - {error_msg}")
                traceback.print_exc()
                results[filename] = {"chunks": 0, "status": "error", "error": error_msg}
                
                # Even if embedding failed, try to move file to processed
                # so user doesn't get stuck re-processing the same file
                try:
                    dst_path = os.path.join(processed_dir, filename)
                    if os.path.exists(dst_path):
                        os.remove(dst_path)
                    shutil.move(src_path, dst_path)
                    print(f"[Knowledge] ✓ {filename}: moved to processed despite error")
                except Exception as move_err:
                    print(f"[Knowledge] ✗ {filename}: failed to move - {move_err}")
        
        print(f"[Knowledge] Processing complete. Results: {results}")
        return {"status": "success", "results": results}
    except Exception as e:
        print(f"[Knowledge] Fatal error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
